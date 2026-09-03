from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\fleetos\deliverables\digital-signature-feature-plan.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5C6570"
WHITE = "FFFFFF"


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=NAVY, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    set_font(run, size=size, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_font(run, size=11, color=NAVY, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_para(doc, text="", size=11, color=NAVY, bold=False, after=6, before=0):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(15 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, color=BLUE if level == 1 else NAVY, bold=True)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    set_rtl(p2)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    r2 = p2.add_run(text)
    set_font(r2, size=10.5, color=NAVY)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=NAVY, size=10.5)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell_text(cells[index], value, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FleetOS | אפיון פיצ'ר טפסים לחתימה דיגיטלית")
    set_font(r, size=8.5, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    set_rtl(title)
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("אפיון פיצ'ר: טפסים לחתימה דיגיטלית")
    set_font(title_run, size=23, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    set_rtl(subtitle)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("מסמך עבודה לפיתוח ב-FleetOS | גרסת תכנון 1.3")
    set_font(subtitle_run, size=11.5, color=MUTED)

    add_table(doc, ["נושא", "פרטים"], [
        ("מטרת המסמך", "הנחיות מוצר, UX, נתונים ואבטחה להטמעת טפסים דיגיטליים לנהגים."),
        ("קהל יעד", "מפתחי Frontend, Backend, Supabase ו-QA."),
        ("החלטת מוצר", "לכל נהג ולכל סוג טופס יש עד מסמך signed אחד ועד מסמך pending אחד."),
        ("תאריך", "26/08/2026"),
    ], [1.55, 4.95])

    add_callout(doc, "החלטה מחייבת", "כאשר הנהג חותם על טופס חדש מאותו סוג, המסמך החתום הקודם נמחק מהמערכת ומהאחסון הפרטי. עד לחתימה החדשה, המסמך החתום הקודם נשאר זמין לצפייה.")

    add_heading(doc, "1. מטרת הפיצ'ר")
    add_para(doc, "המערכת תאפשר למנהל חברה לשלוח טפסי PDF לנהגים לחתימה דיגיטלית, ולנהג לצפות, לחתום ולהוריד את המסמכים הפעילים שלו. הפיצ'ר מיועד למסמכים תקופתיים כגון הצהרת בריאות, אישור נהלי בטיחות, קבלת רכב, רישיון נהיגה וטפסים מותאמים אישית.")
    add_para(doc, "המסמך החתום הוא תוצר סופי: לאחר החתימה נוצר PDF חתום ונשמר באחסון פרטי. כאשר הנהג חותם על מסמך חדש מאותו סוג, העותק הקודם נמחק.")

    add_heading(doc, "2. כלל עסקי ומחזור חיי מסמך")
    add_table(doc, ["שלב", "פעולת מערכת", "מה המשתמש רואה"], [
        ("תבנית", "המנהל מגדיר PDF, סוג טופס ומיקום חתימה.", "הטופס זמין לשליחה לנהגים."),
        ("שליחה", "נוצרת שליחה חדשה לנהג במצב pending.", "הנהג מקבל התראה ורואה 'חתום עכשיו'."),
        ("חתימה", "השרת יוצר PDF חתום ושומר זמן חתימה.", "הנהג והמנהל רואים מסמך פעיל."),
        ("תזכורת חידוש", "אם האדמין הפעיל תזכורת, הוא מקבל התראה במועד שבחר.", "אין שליחה אוטומטית של טופס חדש."),
        ("החלפה", "הנהג חותם על המסמך החדש.", "החדש הופך למסמך החתום הפעיל והקודם נמחק."),
        ("מחיקה", "נמחקים הרשומה והקבצים מה-Storage.", "המסמך אינו זמין לאף משתמש."),
    ], [1.2, 3.15, 2.15])

    add_heading(doc, "3. הרשאות ותפקידי משתמש")
    add_table(doc, ["פעולה", "נהג", "מנהל חברה"], [
        ("צפייה במסמך", "רק המסמכים הפעילים האישיים שלו.", "רק מסמכים של נהגים בחברה שלו."),
        ("חתימה", "יכול לחתום רק על מסמך pending ששייך לו.", "אינו חותם בשם הנהג."),
        ("הורדה", "יכול להוריד את המסמך החתום שלו.", "יכול להוריד מסמך של נהג בחברה שלו."),
        ("שליחה חדשה", "לא מורשה.", "יכול לשלוח או לחדש טופס לנהג."),
        ("מחיקה", "לא מורשה.", "יכול למחוק מסמך פעיל עם אישור מפורש."),
    ], [1.55, 2.85, 2.1])

    add_heading(doc, "4. מסכים וחוויית משתמש")
    add_table(doc, ["מסך", "משתמש", "תכולה עיקרית"], [
        ("מסמכים לחתימה", "נהג", "רשימת מסמכים pending ו-signed בלבד, תוקף וכפתור 'חתום עכשיו'."),
        ("חתימה על מסמך", "נהג", "פתיחת PDF, קנבס חתימה, ניקוי החתימה ואישור סופי."),
        ("מסמכי נהג", "מנהל", "רשימת מסמכים של נהג אחד, צפייה, הורדה, שליחה חדשה ומחיקה."),
        ("תבניות טפסים", "מנהל", "יצירת תבנית PDF, הגדרת סוג, תוקף ומיקום חתימה."),
    ], [1.55, 1.2, 3.75])
    add_callout(doc, "המלצת UX", "מסך המנהל יוצג במבנה של כרטיס נהג עם סטטוס, תוקף ומספר מסמכים בצד אחד; ולצדו פירוט המסמך הפעיל, פעולות צפייה/הורדה/מחיקה וכפתור שליחה מחדש.")

    add_heading(doc, "5. זרימת נהג")
    add_table(doc, ["צעד", "התנהגות נדרשת"], [
        ("1. קבלת התראה", "הנהג מקבל הודעה על טופס חדש ומנווט לרשימת המסמכים."),
        ("2. פתיחת מסמך", "הנהג פותח את ה-PDF באמצעות קישור זמני ומאובטח."),
        ("3. חתימה", "הנהג מצייר חתימה בקנבס. ניסיון לאשר קנבס ריק נחסם עם הודעה בעברית."),
        ("4. אישור", "המערכת יוצרת ושומרת PDF חתום. במקרה כשל, החתימה נשארת במסך לצורך ניסיון חוזר."),
        ("5. צפייה עתידית", "לאחר חתימה, המסמך נגיש לצפייה ולהורדה עד שיוחלף בטופס חדש."),
    ], [1.35, 5.15])

    add_heading(doc, "6. כללים משלימים וסגירת מקרי קצה")
    add_table(doc, ["נושא", "החלטה מחייבת"], [
        ("מסמך ישן ומסמך חדש", "כשיש טופס pending חדש, המסמך signed הישן נשאר מוצג תחת 'מסמך פעיל'. החדש מוצג בנפרד תחת 'ממתין לחתימה'. לאחר החתימה החדש נשאר והישן נמחק."),
        ("שליחה כפולה", "מותר pending אחד בלבד לכל נהג ולכל סוג טופס. שליחה נוספת לפני חתימה מוחקת רק את ה-pending הקודם ומחליפה אותו; המסמך signed נשאר ללא שינוי."),
        ("התראת חידוש", "ברירת המחדל היא כבויה. אם האדמין מפעיל אותה ולא קובע תאריך, מועד ההתראה הוא 12 חודשים מ-signed_at לפי שעון ישראל. האדמין רשאי לבחור תאריך אחר בעת השליחה."),
        ("התנהגות במועד ההתראה", "האדמין מקבל התראה ובוחר אם לשלוח טופס חדש. המערכת אינה שולחת טופס אוטומטית ואינה מוסיפה סטטוס חדש."),
        ("שליחה ראשונה", "בגרסת MVP מנהל שולח לנהג בודד בלבד. שליחה לקבוצת נהגים תתווסף בשלב עתידי."),
        ("מחיקה במקביל", "אם מנהל מוחק pending בזמן שהנהג חותם, השרת דוחה את החתימה בהודעה שהטופס אינו זמין עוד."),
        ("תבניות", "תבנית שכבר נשלחה אינה נערכת או נמחקת. ניתן להשבית אותה וליצור תבנית חדשה."),
        ("אישור הנהג", "לפני 'אשר וחתום' הנהג מסמן: 'קראתי את המסמך ואני מאשר/ת את תוכנו'. בלי האישור והחתימה הפעולה נחסמת."),
    ], [1.75, 4.75])

    add_heading(doc, "7. כללי אמינות, אבטחה ומקרי קצה")
    add_table(doc, ["נושא", "החלטה מחייבת"], [
        ("קישור צפייה אחרי מחיקה", "קישור צפייה חתום תקף ל-5 דקות. לאחר מחיקה לא נוצרים קישורים חדשים, אך קישור שכבר נפתח עשוי לעבוד עד לפקיעתו."),
        ("פעולת חתימה חוזרת", "לכל אישור חתימה נשלח מזהה פעולה ייחודי. לחיצה כפולה או ניסיון חוזר לאחר ניתוק מחזירים את אותה תוצאה ולא יוצרים מסמך נוסף."),
        ("פעולות מקבילות", "השרת נועל את הטופס בזמן חתימה, מחיקה או החלפה. הפעולה הראשונה שמושלמת קובעת; השנייה נדחית עם הודעה שהטופס אינו זמין."),
        ("עזיבת נהג את החברה", "כאשר נהג מושבת, נמחק או מועבר לחברה אחרת, הגישה שלו למסמכי החברה נחסמת מיד וכל הטפסים pending שלו מבוטלים ונמחקים. מסמכים signed נשארים זמינים למנהל בלבד עד למחיקה ידנית או החלפה."),
        ("מגבלות קובץ", "בגרסה הראשונה מתקבלים PDF בלבד, עד 10MB, עד 10 עמודים, ללא סיסמה, ועם שדה חתימת נהג אחד בלבד. קובץ שלא עומד בתנאים נדחה לפני שמירת התבנית."),
        ("עמודי PDF מסובבים", "מיקום החתימה נשמר ביחס לעמוד ומתחשב בכיוון, בסיבוב ובמידות העמוד. השרת בודק שהמלבן כולו נמצא בתוך גבולות העמוד."),
        ("סגירת אפליקציה", "לפני אישור אין שמירה. אם הרשת נכשלת לאחר אישור, החתימה נשארת במסך כל עוד האפליקציה פתוחה. סגירת האפליקציה מוחקת חתימה שלא נשלחה."),
        ("מחיקה ידנית עם pending", "מחיקת מסמך signed ישן אינה מוחקת טופס pending חדש. הנהג נשאר עם הטופס הממתין עד שיחתום או שהמנהל ימחק גם אותו."),
        ("עותק תבנית", "בעת השליחה נשמרים עותק PDF, גרסת תבנית ומיקום החתימה. יצירת ה-PDF החתום משתמשת תמיד בעותק זה, גם אם התבנית השתנתה או הושבתה."),
    ], [1.75, 4.75])

    add_heading(doc, "8. זרימת מנהל")
    add_table(doc, ["צעד", "התנהגות נדרשת"], [
        ("1. בחירת נהג", "המנהל פותח את אזור 'טפסים דיגיטליים' מתוך פרטי הנהג."),
        ("2. צפייה בסטטוס", "מוצגים סוג הטופס, סטטוס, תאריך שליחה, תאריך חתימה ומועד התראת חידוש אם הופעלה."),
        ("3. שליחה חדשה", "המנהל בוחר תבנית ושולח אותה לנהג. המסמך הקודם נשאר זמין עד שהנהג חותם על החדש."),
        ("4. מחיקה ידנית", "נדרשת חלונית אישור שמבהירה שהמחיקה בלתי הפיכה ומוחקת גם את הקובץ."),
        ("5. הורדה", "המנהל מקבל קישור זמני ל-PDF החתום בלבד, לא למסמך מקור לא חתום."),
    ], [1.35, 5.15])

    add_heading(doc, "9. מיקום החתימה על ה-PDF")
    add_para(doc, "המנהל הוא שקובע את מיקום החתימה פעם אחת, בעת יצירת תבנית הטופס. הנהג אינו בוחר מיקום ואינו מציב את החתימה ידנית על גבי המסמך.")
    add_table(doc, ["שלב", "התנהגות נדרשת"], [
        ("1. העלאת PDF", "המנהל מעלה את קובץ ה-PDF המקורי של הטופס."),
        ("2. סימון השדה", "המערכת מציגה את עמוד ה-PDF עצמו. המנהל גורר מלבן 'חתימת הנהג' אל השדה המיועד ויכול לשנות את גודלו."),
        ("3. שמירת המיקום", "נשמרים מספר העמוד, מיקום X/Y ורוחב/גובה יחסיים לעמוד."),
        ("4. חתימת הנהג", "הנהג מצייר בקנבס. השרת ממיר את החתימה לתמונה שקופה ומטביע אותה בדיוק בתוך המלבן שהוגדר."),
        ("5. PDF סופי", "נוצר PDF חתום חדש לצפייה ולהורדה. ה-PDF המקורי לעולם אינו משתנה."),
    ], [1.4, 5.1])
    add_callout(doc, "כלל UX", "חובה להציג למנהל את עמוד ה-PDF האמיתי בזמן בחירת המיקום, ולא תרשים כללי של עמוד. כך החתימה תמוקם בדיוק בתוך השדה הנכון בטופס.")

    add_heading(doc, "10. מודל נתונים מוצע")
    add_table(doc, ["ישות", "שדות עיקריים", "ייעוד"], [
        ("document_templates", "company_id, title, file_path, signature coordinates", "תבניות PDF קבועות של החברה."),
        ("driver_document_sends", "driver_id, template_id, template_version, template_snapshot_path, status, signed_file_path, sent_at, signed_at, renewal_reminder_at, renewal_reminder_sent_at", "המופע הפעיל או הממתין של טופס לנהג."),
    ], [1.55, 3.45, 1.5])
    add_para(doc, "הסטטוסים היחידים ב-driver_document_sends הם pending ו-signed. אין סטטוס replaced ואין היסטוריית החלפות. נדרשת מגבלה בשרת: לכל נהג ולכל תבנית לכל היותר pending אחד ו-signed אחד. אין לערבב את הפיצ'ר עם טבלת documents הקיימת, המיועדת למסמכי נהג ורכב כלליים.", size=10.5, color=MUTED, after=4)

    add_heading(doc, "11. Storage, אבטחה והרשאות")
    add_table(doc, ["נושא", "דרישה"], [
        ("אחסון", "Bucket פרטי נפרד, לדוגמה signed-driver-documents."),
        ("גישה לקבצים", "קישורים זמניים בלבד (signed URLs). אין קישורים ציבוריים."),
        ("RLS לנהג", "נהג רואה רק שורות שבהן driver_id שווה למשתמש המחובר."),
        ("RLS למנהל", "מנהל רואה נהגים ומסמכים של החברה שלו בלבד."),
        ("פעולות רגישות", "מחיקה, בדיקת התראות חידוש, יצירת PDF חתום וקבלת URL מבוצעות דרך Edge Functions."),
        ("Service Role", "אסור לחשוף Service Role באפליקציית Expo. השימוש בו, אם נדרש, רק בצד השרת."),
    ], [1.55, 4.95])

    add_heading(doc, "12. Edge Functions והתראות חידוש")
    add_table(doc, ["Function", "אחריות"], [
        ("send-driver-document", "יצירת שליחה חדשה, אימות הרשאת מנהל ושמירת snapshot של התבנית."),
        ("sign-driver-document", "אימות בעלות הנהג, עיבוד החתימה, יצירת PDF חתום, עדכון סטטוס ומחיקת המסמך החתום הקודם."),
        ("delete-driver-document", "מחיקה מבוקרת של הרשומה ושל קבצי ה-Storage."),
        ("check-document-renewal-reminders", "משימה יומית שמאתרת התראות חידוש שהגיע מועדן, שולחת התראה אחת לאדמין ולא יוצרת טופס חדש."),
        ("get-signed-document-url", "בדיקת הרשאה והנפקת signed URL קצר תוקף למסמך החתום."),
    ], [2.25, 4.25])
    add_para(doc, "החלפת מסמך מתבצעת בשרת רק לאחר שהחתימה החדשה נשמרה בהצלחה: נעילת הרשומה הפעילה, יצירת PDF חדש, עדכון המסמך החדש ל-signed ומחיקת הקבצים הקודמים. יש להוסיף מנגנון retry במקרה שבו מחיקת Storage נכשלת לאחר עדכון מסד הנתונים.", size=10.5)

    add_heading(doc, "13. ניווט ושילוב בקוד הקיים")
    add_table(doc, ["מיקום קיים", "שינוי מוצע"], [
        ("DriverDetailScreen", "להוסיף כניסה ל'טפסים דיגיטליים' עבור הנהג שנבחר."),
        ("DriverDocuments", "להוסיף אזור נפרד למסמכים חתומים/ממתינים, בלי לשנות מסמכי נהג כלליים."),
        ("navigation/types.ts", "להוסיף Routes למסך מסמכי נהג, מסך חתימה ומסך מסמכי נהג עבור מנהל."),
        ("lib/documents.ts", "לשמור את דפוס signed URL והורדה, אך ליצור service נפרד לטפסים חתומים."),
    ], [2.0, 4.5])

    add_heading(doc, "14. בדיקות קבלה")
    add_table(doc, ["מספר", "בדיקה"], [
        ("1", "נהג אינו יכול לצפות, לחתום או להוריד מסמך של נהג אחר."),
        ("2", "מנהל אינו יכול לצפות או למחוק מסמך של חברה אחרת."),
        ("3", "חתימה ריקה נחסמת ולא משנה את מצב המסמך."),
        ("4", "לאחר חתימה נוצר PDF חתום והסטטוס הופך ל-signed."),
        ("5", "שליחת טופס חדש אינה מוחקת את המסמך החתום הקודם; החתימה על החדש מוחקת אותו מה-DB ומה-Storage."),
        ("6", "מחיקה ידנית של מנהל מוחקת את הרשומה, הקבצים ומונעת גישה דרך URL חדש."),
        ("7", "התראת חידוש כבויה אינה יוצרת התראה או טופס חדש; התראה פעילה שולחת הודעה אחת לאדמין בלבד."),
        ("8", "הזרימה עובדת ב-iOS וב-Android, כולל ציור חתימה באצבע."),
        ("9", "שליחה נוספת בזמן pending מחליפה רק את ה-pending הקודם ולא מוחקת מסמך signed פעיל."),
        ("10", "חתימה ללא סימון 'קראתי את המסמך' נחסמת."),
        ("11", "אם אדמין מפעיל התראת חידוש בלי תאריך, היא נשלחת שנה מ-signed_at לפי שעון ישראל."),
        ("12", "קישור צפייה שכבר נוצר תקף לכל היותר 5 דקות לאחר מחיקת המסמך."),
        ("13", "לחיצה כפולה על חתימה או ניתוק רשת לאחר חתימה אינם יוצרים PDF כפול."),
        ("14", "נהג שהוסר מהחברה אינו יכול לפתוח או לחתום על מסמכי החברה."),
        ("15", "PDF מסובב מטביע חתימה במקום המדויק שהוגדר, בתוך גבולות העמוד."),
    ], [0.7, 5.8])

    add_heading(doc, "15. סדר פיתוח מומלץ")
    add_table(doc, ["שלב", "תוצרים"], [
        ("שלב א - נתונים ואבטחה", "מיגרציות, Buckets, RLS, מגבלות pending/signed ו-Edge Functions בסיסיות."),
        ("שלב ב - מנהל", "תבניות, שליחה לנהג, רשימת סטטוסים, צפייה והורדה."),
        ("שלב ג - נהג", "רשימת מסמכים, מסך חתימה, יצירת PDF חתום והתראות."),
        ("שלב ד - התראות חידוש ומחיקה", "בדיקה יומית של תזכורות שהופעלו, החלפה לאחר חתימה, מחיקה מ-Storage ו-Retry."),
        ("שלב ה - איכות", "בדיקות RLS, Unit Tests, בדיקות Edge Functions ובדיקות ידניות במכשירי iPhone ו-Android."),
    ], [1.65, 4.85])

    add_callout(doc, "החלטה סופית", "למערכת יש שני סטטוסים בלבד: pending ו-signed. העותק החתום הקודם נשאר זמין עד שהחדש נחתם, ואז נמחק לצמיתות. אין סטטוס replaced, אין ארכיון ואין היסטוריית החלפות במוצר.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
